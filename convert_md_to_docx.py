#!/usr/bin/env python3
"""
Simple markdown to Word document converter for the research paper.
Handles basic markdown formatting including headers, paragraphs, bold, italic, lists, and links.
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_inline_formatting(text, paragraph):
    """Parse inline markdown formatting (bold, italic, links) and add runs to paragraph."""
    # Pattern to match **bold**, *italic*, and [link](url)
    pattern = r'\*\*(.+?)\*\*|\*(.+?)\*|\[(.+?)\]\((.+?)\)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Add formatted text
        if match.group(1):  # Bold
            run = paragraph.add_run(match.group(1))
            run.bold = True
        elif match.group(2):  # Italic
            run = paragraph.add_run(match.group(2))
            run.italic = True
        elif match.group(3):  # Link
            run = paragraph.add_run(match.group(3))
            run.font.color.rgb = None  # Keep default color
            # Note: python-docx doesn't easily support hyperlinks, so we just show the text

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_list = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if line.startswith('---'):
            i += 1
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                # Title - centered and larger
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
            elif level == 2:
                # Main section
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
            elif level == 3:
                # Subsection
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(12)

        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(text, p)
            in_list = True

        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(text, p)
            in_list = True

        # Handle empty lines
        elif not line.strip():
            if in_list:
                in_list = False
            else:
                doc.add_paragraph()

        # Handle regular paragraphs
        else:
            # Check if it's an author name or course info (centered text)
            if i < 10 and not line.startswith('#') and line.strip():
                p = doc.add_paragraph()
                if 'Computer Architecture' in line or 'RPN1' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                parse_inline_formatting(line, p)
            else:
                p = doc.add_paragraph()
                parse_inline_formatting(line, p)
            in_list = False

        i += 1

    # Save the document
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == '__main__':
    convert_markdown_to_docx(
        '/home/user/paper/research_paper_improved.md',
        '/home/user/paper/research_paper_improved_new.docx'
    )
